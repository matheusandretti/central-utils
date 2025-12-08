# api/gerador_atas_core.py
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any

from docx import Document
from docx.shared import Pt

# =========================
# Diretórios de trabalho
# =========================

BASE_DIR = Path(__file__).resolve().parent.parent  # raiz do projeto Python
DATA_DIR = BASE_DIR / "data"
MODELOS_DIR = DATA_DIR / "atas_modelos"
SAIDA_DIR = DATA_DIR / "atas_geradas"

MODELOS_DIR.mkdir(parents=True, exist_ok=True)
SAIDA_DIR.mkdir(parents=True, exist_ok=True)

# =========================
# Utilidades de formatação
# =========================


def apenas_digitos(texto: str) -> str:
    return "".join(ch for ch in texto if ch.isdigit())


def formatar_cnpj(texto: str) -> str:
    digitos = apenas_digitos(texto)
    if len(digitos) != 14:
        return texto
    return (
        f"{digitos[0:2]}.{digitos[2:5]}.{digitos[5:8]}/{digitos[8:12]}-{digitos[12:14]}"
    )


def formatar_cpf(texto: str) -> str:
    digitos = apenas_digitos(texto)
    if len(digitos) != 11:
        return texto
    return f"{digitos[0:3]}.{digitos[3:6]}.{digitos[6:9]}-{digitos[9:11]}"


def formatar_cep(texto: str) -> str:
    digitos = apenas_digitos(texto)
    if len(digitos) != 8:
        return texto
    return f"{digitos[0:5]}-{digitos[5:8]}"


def formatar_cidade(texto: str) -> str:
    return " ".join(p.capitalize() for p in texto.strip().split())


def normalizar_estado(texto: str) -> str:
    letras = [c for c in texto.upper() if c.isalpha()]
    if len(letras) != 2:
        raise ValueError("ESTADO precisa ter exatamente 2 letras.")
    return "".join(letras)


def normalizar_data(texto: str) -> str:
    txt = texto.strip()
    if not txt:
        raise ValueError("Data vazia")

    txt2 = txt.replace("-", "/")
    try:
        dt = datetime.strptime(txt2, "%d/%m/%Y")
        return dt.strftime("%d/%m/%Y")
    except ValueError:
        pass

    dig = apenas_digitos(txt)
    if len(dig) == 8:
        dia = int(dig[0:2])
        mes = int(dig[2:4])
        ano = int(dig[4:8])
        dt = datetime(ano, mes, dia)
        return dt.strftime("%d/%m/%Y")

    raise ValueError(f"Formato de data inválido: {texto}")


def converter_moeda_para_centavos(texto: str) -> int:
    txt = texto.strip()
    if not txt:
        return 0
    sinal = -1 if txt.startswith("-") else 1
    digitos = apenas_digitos(txt)
    if not digitos:
        return 0
    return sinal * int(digitos)


def formatar_moeda_de_centavos(centavos: int) -> str:
    sinal = "-" if centavos < 0 else ""
    centavos = abs(centavos)
    inteiro = centavos // 100
    frac = centavos % 100

    partes = []
    s = str(inteiro)
    while len(s) > 3:
        partes.insert(0, s[-3:])
        s = s[:-3]
    partes.insert(0, s)
    inteiro_fmt = ".".join(partes)

    return f"{sinal}{inteiro_fmt},{frac:02d}"


DIAS_EXTENSO = {
    1: "UM",
    2: "DOIS",
    3: "TRÊS",
    4: "QUATRO",
    5: "CINCO",
    6: "SEIS",
    7: "SETE",
    8: "OITO",
    9: "NOVE",
    10: "DEZ",
    11: "ONZE",
    12: "DOZE",
    13: "TREZE",
    14: "QUATORZE",
    15: "QUINZE",
    16: "DEZESSEIS",
    17: "DEZESSETE",
    18: "DEZOITO",
    19: "DEZENOVE",
    20: "VINTE",
    21: "VINTE E UM",
    22: "VINTE E DOIS",
    23: "VINTE E TRÊS",
    24: "VINTE E QUATRO",
    25: "VINTE E CINCO",
    26: "VINTE E SEIS",
    27: "VINTE E SETE",
    28: "VINTE E OITO",
    29: "VINTE E NOVE",
    30: "TRINTA",
    31: "TRINTA E UM",
}

MESES_EXTENSO = {
    1: "JANEIRO",
    2: "FEVEREIRO",
    3: "MARÇO",
    4: "ABRIL",
    5: "MAIO",
    6: "JUNHO",
    7: "JULHO",
    8: "AGOSTO",
    9: "SETEMBRO",
    10: "OUTUBRO",
    11: "NOVEMBRO",
    12: "DEZEMBRO",
}

CAMPOS_DERIVADOS = {"DIA(S)", "MÊS"}


def guess_tipo_campo(nome: str) -> str:
    if nome == "CNPJ":
        return "cnpj"
    if nome == "CEP":
        return "cep"
    if "CPF" in nome:
        return "cpf"
    if nome.startswith("LUCRO_"):
        return "money"
    if nome == "DATA":
        return "data"
    return "texto"


def label_amigavel(nome: str) -> str:
    especiais = {
        "NIRE": "NIRE",
        "CNPJ": "CNPJ",
        "CEP": "CEP",
        "DATA": "Data (DD/MM/AAAA)",
        "CIDADE": "Cidade",
        "ESTADO": "Estado (UF)",
        "NUMERO_RUA": "Número",
        "LUCRO_2024": "Lucro 2024",
        "LUCRO_2023": "Lucro 2023",
        "LUCRO_2022": "Lucro 2022",
        "LUCRO_2021_E_OUTROS": "Lucro 2021 e outros",
        "LUCRO_TOTAL": "Lucro total",
    }
    if nome in especiais:
        return especiais[nome]
    return nome.replace("_", " ").title()


# =========================
# Placeholders e DOCX
# =========================


def encontrar_placeholders_docx(model_path: Path) -> List[str]:
    doc = Document(str(model_path))
    encontrados = set()

    def scan_paragraphs(paragraphs):
        for p in paragraphs:
            texto = "".join(run.text for run in p.runs)
            import re

            for m in re.finditer(r"\{\{([^}]+)\}\}", texto):
                encontrados.add(m.group(1))

    scan_paragraphs(doc.paragraphs)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                scan_paragraphs(celula.paragraphs)

    return sorted(encontrados)


def substituir_placeholders_no_texto(texto: str, dados: Dict[str, str]) -> str:
    if not texto:
        return texto
    resultado = texto
    for campo, valor in dados.items():
        placeholder = "{{" + campo + "}}"
        resultado = resultado.replace(placeholder, valor)
    return resultado


def processar_paragrafos_substituicao(paragraphs, dados: Dict[str, str]):
    """
    Substitui {{...}} preservando fonte/tamanho do parágrafo
    e aplica as formatações especiais:
    - título "ATA DE..." em negrito
    - nome da empresa em negrito
    - linhas que começam com CNPJ em negrito
    - linha do NIRE em maiúsculo, negrito e tamanho 14 (Calibri)
    - linhas que começam com IOB em negrito
    """
    for p in paragraphs:
        # 1) Substituição dos placeholders preservando fonte/tamanho
        texto_original = "".join(run.text for run in p.runs)
        texto_novo = substituir_placeholders_no_texto(texto_original, dados)

        if texto_novo != texto_original:
            first_run = p.runs[0] if p.runs else None
            font_name = (
                first_run.font.name if first_run and first_run.font.name else None
            )
            font_size = (
                first_run.font.size if first_run and first_run.font.size else None
            )

            p.text = texto_novo  # substitui o conteúdo inteiro

            # reaplica fonte/tamanho do primeiro run
            if font_name or font_size:
                for run in p.runs:
                    if font_name:
                        run.font.name = font_name
                    if font_size:
                        run.font.size = font_size

        # 2) Regras de negrito/cabeçalho
        texto = p.text.strip()
        if not texto:
            continue

        texto_upper = texto.upper()

        # 1) TÍTULO: linhas que começam com "ATA DE"
        if texto_upper.startswith("ATA DE"):
            for run in p.runs:
                run.bold = True
            continue

        # 2) NOME DA EMPRESA
        nome_emp = (dados.get("NOME_EMPRESA") or "").strip()
        if nome_emp and texto_upper == nome_emp.upper():
            for run in p.runs:
                run.bold = True
            continue

        # 3) CNPJ
        if texto_upper.startswith("CNPJ"):
            for run in p.runs:
                run.bold = True
            continue

        # 4) NIRE
        nire_val = (dados.get("NIRE") or "").strip().upper()
        if nire_val and (texto_upper == nire_val or texto_upper.startswith("NIRE")):
            p.text = p.text.upper()
            for run in p.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(14)
            continue

        # 5) IOB
        if texto_upper.startswith("IOB"):
            for run in p.runs:
                run.bold = True
            continue


def preencher_documento(modelo_path: Path, dados: Dict[str, str]) -> Document:
    doc = Document(str(modelo_path))
    processar_paragrafos_substituicao(doc.paragraphs, dados)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                processar_paragrafos_substituicao(celula.paragraphs, dados)
    return doc


def montar_nome_arquivo_saida(modelo_path: Path, dados: Dict[str, str]) -> str:
    base = modelo_path.stem
    nome_empresa = dados.get("NOME_EMPRESA") or dados.get("NOME EMPRESA") or ""
    if nome_empresa:
        import re

        safe_nome = re.sub(r'[\\/:*?"<>|]', "_", nome_empresa.strip())
        safe_nome = safe_nome.replace(" ", "_")
        return f"{safe_nome} - ({base}).docx"
    return f"{base}.docx"


# =========================
# Lógica de lucros
# =========================


def calcular_campos_lucros(lucros: List[Dict[str, Any]]) -> Dict[str, str]:
    """
    lucros: lista de dicts {ano: int, valor: str (moeda brasileira)}
    Retorna dict com chaves LUCRO_<ano>, LUCRO_TOTAL, LUCRO_2021_E_OUTROS.
    """
    itens = []
    for item in lucros:
        ano = int(item.get("ano"))
        valor_txt = str(item.get("valor", ""))
        centavos = converter_moeda_para_centavos(valor_txt)
        itens.append({"ano": ano, "centavos": centavos})

    resultado: Dict[str, int] = {}
    total = sum(i["centavos"] for i in itens)
    resultado["LUCRO_TOTAL"] = total

    soma_2021_e_antes = sum(i["centavos"] for i in itens if i["ano"] <= 2021)
    resultado["LUCRO_2021_E_OUTROS"] = soma_2021_e_antes

    for item in itens:
        ano = item["ano"]
        key = f"LUCRO_{ano}"
        resultado[key] = resultado.get(key, 0) + item["centavos"]

    # converte para texto
    return {k: formatar_moeda_de_centavos(v) for k, v in resultado.items()}

def completar_lucros_zerados(modelo_path: Path, dados: Dict[str, str]) -> None:
    """
    Para todo placeholder LUCRO_... encontrado no modelo, se não houver valor
    correspondente em 'dados', preenche com '0,00' para evitar {{LUCRO_XXXX}} visível.
    """
    try:
        placeholders = encontrar_placeholders_docx(modelo_path)
    except Exception:
        return

    for ph in placeholders:
        if ph.startswith("LUCRO_") and ph not in dados:
            dados[ph] = "0,00"

# =========================
# Assinaturas
# =========================


def aplicar_assinaturas_no_doc(
    doc: Document,
    assinaturas_pf: List[Dict[str, str]],
    assinaturas_pj: List[Dict[str, str]],
) -> None:
    if not assinaturas_pf and not assinaturas_pj:
        return

    def add_line(text, bold=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        if bold:
            run.bold = True

    primeira_linha = True

    # PF
    for s in assinaturas_pf:
        nome = (s.get("nome") or "").strip()
        cpf = (s.get("cpf") or "").strip()
        qualif = (s.get("qualificacao") or "").strip()
        if not nome and not cpf and not qualif:
            continue

        if not primeira_linha:
            add_line("", space_before=4, space_after=0)
        primeira_linha = False

        add_line(nome.upper(), bold=True)
        if qualif:
            add_line(qualif)
        if cpf:
            add_line(f"CPF nº: {formatar_cpf(cpf)}")

    # PJ
    for s in assinaturas_pj:
        pj = (s.get("pj") or "").strip()
        rep = (s.get("representante") or "").strip()
        cpf = (s.get("cpf") or "").strip()
        qualif = (s.get("qualificacao") or "").strip()
        if not pj and not rep and not cpf and not qualif:
            continue

        if not primeira_linha:
            add_line("", space_before=4, space_after=0)
        primeira_linha = False

        add_line(pj.upper(), bold=True)
        if rep:
            add_line("Representada por")
            add_line(rep.upper(), bold=True)
        if qualif:
            add_line(qualif)
        if cpf:
            add_line(f"CPF nº: {formatar_cpf(cpf)}")


# =========================
# API de alto nível
# =========================


def listar_modelos() -> List[Dict[str, str]]:
    modelos = []
    for p in sorted(MODELOS_DIR.glob("*.docx")):
        modelos.append({"id": p.name, "fileName": p.name, "displayName": p.stem})
    return modelos


def obter_campos_modelo(modelo_id: str) -> Dict[str, Any]:
    modelo_path = MODELOS_DIR / modelo_id
    if not modelo_path.is_file():
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_id}")

    placeholders = encontrar_placeholders_docx(modelo_path)
    campos = []
    lucros_placeholders = []

    for nome in placeholders:
        if nome.startswith("LUCRO_"):
            lucros_placeholders.append(nome)
            continue
        if nome in CAMPOS_DERIVADOS:
            continue
        campos.append(
            {
                "name": nome,
                "label": label_amigavel(nome),
                "tipo": guess_tipo_campo(nome),
            }
        )

    # ordena para ficar amigável
    def key_func(c):
        nome = c["name"]
        ordem_basica = [
            "NOME_EMPRESA",
            "CNPJ",
            "NIRE",
            "DATA",
            "CIDADE",
            "ESTADO",
            "RUA",
            "NUMERO_RUA",
            "BAIRRO",
            "CEP",
        ]
        if nome in ordem_basica:
            return (0, ordem_basica.index(nome))
        return (1, nome)

    campos.sort(key=key_func)

    return {"campos": campos, "lucrosPlaceholders": lucros_placeholders}


def gerar_ata(
    modelo_id: str,
    campos: Dict[str, str],
    lucros: List[Dict[str, Any]],
    assinaturas_pf: List[Dict[str, str]],
    assinaturas_pj: List[Dict[str, str]],
) -> str:
    modelo_path = MODELOS_DIR / modelo_id
    if not modelo_path.is_file():
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_id}")

    dados = dict(campos)

    # equivalência NOME_EMPRESA / NOME EMPRESA
    if "NOME_EMPRESA" in dados and "NOME EMPRESA" not in dados:
        dados["NOME EMPRESA"] = dados["NOME_EMPRESA"]

    # normalização de campos conhecidos
    if "DATA" in dados:
        dados["DATA"] = normalizar_data(dados["DATA"])
        try:
            dt = datetime.strptime(dados["DATA"], "%d/%m/%Y")
            dia = dt.day
            mes = dt.month
            dados["DIA(S)"] = DIAS_EXTENSO.get(dia, str(dia)).lower()
            dados["MÊS"] = MESES_EXTENSO.get(mes, str(mes)).lower()
        except ValueError:
            pass

    if "CIDADE" in dados:
        dados["CIDADE"] = formatar_cidade(dados["CIDADE"])

    if "ESTADO" in dados:
        dados["ESTADO"] = normalizar_estado(dados["ESTADO"])

    if "CEP" in dados:
        dados["CEP"] = formatar_cep(dados["CEP"])

    if "CNPJ" in dados:
        dados["CNPJ"] = formatar_cnpj(dados["CNPJ"])

    # formata quaisquer campos com "CPF" no nome
    for k, v in list(dados.items()):
        if "CPF" in k:
            dados[k] = formatar_cpf(v)

    # lucros
    if lucros:
        dados.update(calcular_campos_lucros(lucros))
        
    # Preenche lucros que não vieram da tela com 0,00
    completar_lucros_zerados(modelo_path, dados)

    # gerar documento
    doc = preencher_documento(modelo_path, dados)
    aplicar_assinaturas_no_doc(doc, assinaturas_pf, assinaturas_pj)

    nome_saida = montar_nome_arquivo_saida(modelo_path, dados)
    caminho_saida = SAIDA_DIR / nome_saida
    SAIDA_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(caminho_saida))

    return nome_saida
